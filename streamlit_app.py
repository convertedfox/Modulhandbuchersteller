import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from typing import Dict, Any

def add_formatted_section(document: Document, data: Dict[str, Any], alter_studiengang: str) -> str:
    if alter_studiengang != data["Studiengang"]:
        studiengang = document.add_heading(data["Studiengang"], level=0)
        studiengang.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        studiengang.runs[0].font.size = Pt(22)
        try:
            studiengang.style = document.styles["Header_Studiengang"]
        except KeyError:
            pass
        alter_studiengang = data["Studiengang"]
    
    # Modultitel
    title = document.add_heading(str(data['Modultitel']), level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title.runs[0].font.size = Pt(20)


    subtitle = document.add_paragraph(f"{data['Modulcode']} - ECTS: {data['Credits']}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(11)
    # Fügen Sie eine Leerzeile hinzu
    document.add_paragraph()

    for section in ['Kompetenzbeschreibung – Kurzform', 'Lehrinhalte']:
        heading = document.add_paragraph()
        heading_run = heading.add_run(f"{section}:")
        font = heading_run.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.italic = True
        font.color.rgb = RGBColor(255, 0, 0)  # Rot  
        heading_run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        document.add_paragraph() #Leerzeile
        inhalt = document.add_paragraph()
        inhalt_run = inhalt.add_run(str(data[section]))
        font = inhalt_run.font
        font.size = Pt(10)
        document.add_paragraph()
  
    
    link_heading = document.add_paragraph()
    link_heading_run = link_heading.add_run("Internet-Link mit Detailbeschreibung und Terminen:")
    font = link_heading_run.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.italic = True
    font.color.rgb = RGBColor(255, 0, 0)  # Rot
    link_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph()

    link = document.add_paragraph()
    link.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hyperlink = link.add_run(str(data['Link']))
    hyperlink.font.underline = True

    document.add_page_break()
    return alter_studiengang

def excel_to_dict(uploaded_excel: io.BytesIO) -> list:
    df = pd.read_excel(uploaded_excel, sheet_name=0)
    return df.to_dict(orient='records')

def dict_to_word(uploaded_template: io.BytesIO, uploaded_excel: io.BytesIO) -> tuple:
    data = excel_to_dict(uploaded_excel)
    file_name = f"Ausgabe_Modulhandbuch_{datetime.now().strftime('%Y%m%d_%H%M')}"
    doc = Document(uploaded_template)
    doc.add_page_break()
    alter_studiengang = ""
    for module_data in data:
        alter_studiengang = add_formatted_section(doc, module_data, alter_studiengang)
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes, file_name

def main():
    st.image('./dhbw_cas_logo.jpg')
    st.title("Modulhandbuchersteller für das DHBW CAS mit Vorlagenunterstützung")

    st.markdown("""
    ## Anleitung
    1. Laden Sie eine Excel-Datei mit den Moduldaten hoch.
    2. Laden Sie eine Word-Vorlage für das Modulhandbuch hoch.
    3. Klicken Sie auf 'Konvertieren', um das Modulhandbuch zu erstellen.
    
    **Hinweis:** Die Excel-Datei sollte folgende Spalten in dieser Reihenfolge und mit dieser Benamung enthalten: Studiengang, Modultitel, Modulcode, Credits, Kompetenzbeschreibung – Kurzform, Lehrinhalte, Link
    """)

    uploaded_excel = st.file_uploader("Wählen Sie eine Excel-Datei aus", type="xlsx")
    uploaded_template = st.file_uploader("Wählen Sie eine Word-Vorlage aus", type="docx")

    if uploaded_excel and uploaded_template:
        if st.button("Konvertieren"):
            try:
                doc_bytes, file_name = dict_to_word(uploaded_template, uploaded_excel)
                st.success(f"Konvertierung abgeschlossen. Die Word-Datei heißt '{file_name}.docx'.")
                st.download_button(
                    label="Word-Datei herunterladen",
                    data=doc_bytes,
                    file_name=f"{file_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Ein Fehler ist aufgetreten: {str(e)}")
                st.error("Bitte stellen Sie sicher, dass die hochgeladenen Dateien gültig sind und das erwartete Format haben.")
    elif uploaded_excel or uploaded_template:
        st.warning("Bitte laden Sie sowohl eine Excel-Datei als auch eine Word-Vorlage hoch.")

if __name__ == "__main__":
    main()