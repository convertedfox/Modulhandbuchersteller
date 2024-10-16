import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

alter_studiengang = None

def add_formatted_section(document, data):
    # Studiengang - nur, wenn noch nicht vorhanden
    global alter_studiengang
    if not alter_studiengang == data["Studiengang"]:
        studiengang = document.add_heading(data["Studiengang"], level = 0)
        studiengang.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        studiengang.font_size = Pt(16)
        alter_studiengang = data["Studiengang"]
    
    # Titel
    Titel = str(data['Modultitel'])
    title = document.add_heading(Titel, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title.runs[0]
    title_run.font.size = Pt(14)

    # Untertitel
    subtitle = document.add_paragraph(str(data['Modulcode']) + " - ECTS: " + str(data["Credits"]))
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    # Kompetenzbeschreibung
    comp_heading = document.add_paragraph()
    comp_heading.add_run("Kompetenzbeschreibung – Kurzform:").bold = True
    comp_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    comp_para = document.add_paragraph(str(data['Kompetenzbeschreibung – Kurzform']))
    comp_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Lehrinhalte
    content_heading = document.add_paragraph()
    content_heading.add_run("Lehrinhalte:").bold = True
    content_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    content_item = document.add_paragraph(str(data["Lehrinhalte"]))
    content_item.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Internet-Link
    link_heading = document.add_paragraph()
    link_heading.add_run("Internet-Link mit Detailbeschreibung und Terminen:").bold = True
    link_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    link = document.add_paragraph()
    link.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hyperlink = link.add_run(str(data['Link']))
    hyperlink.font.underline = True
    #hyperlink.font.color.rgb = (0, 0, 255)  # Blau

    document.add_page_break()

def excel_to_dict(uploaded_excel):
    df = pd.read_excel(uploaded_excel, sheet_name=0)  # sheet_name=0 gibt das erste Arbeitsblatt an
    df = df.astype(str) # alles gleich in Strings umwandeln
    dictionary = {} 
    # DataFrame in ein Dictionary umwandeln
    # Das orient='records' sorgt dafür, dass jede Zeile als Dictionary-Eintrag gespeichert wird
    dictionary = df.to_dict(orient='records')  
    return dictionary

def dict_to_word(uploaded_template, uploaded_excel):
    data = excel_to_dict(uploaded_excel)
    file_name = "Ausgabe_Modulhandbuch_" + str(datetime.now().strftime("%Y%m%d_%H%M"))
    doc = Document(uploaded_template)
    # Fügen Sie einen Seitenumbruch am Ende des bestehenden Inhalts hinzu
    doc.add_page_break()
    for module_data in data:
        add_formatted_section(doc, module_data)
    
    # Word-Dokument in Bytes-Objekt speichern
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes, file_name

st.image('./dhbw_cas_logo.jpg')
st.title("Modulhandbuchersteller für das DHBW CAS mit Vorlagenunterstützung")

uploaded_excel = st.file_uploader("Wählen Sie eine Excel-Datei aus", type="xlsx")
uploaded_template = st.file_uploader("Wählen Sie eine Word-Vorlage aus", type="docx")

if uploaded_excel is not None:
    try:
        if st.button("Konvertieren"):
            if uploaded_template:
                doc_bytes, file_name = dict_to_word(uploaded_excel, uploaded_template)
            else:
                doc_bytes, file_name = dict_to_word(uploaded_excel, None)
            
            st.success(f"Konvertierung abgeschlossen. Die Word-Datei heißt '{file_name}.docx'.")
            
            st.download_button(
                label="Word-Datei herunterladen",
                data=doc_bytes,
                file_name=f"{file_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"Ein Fehler ist aufgetreten: {str(e)}")
        st.error("Bitte stellen Sie sicher, dass die hochgeladenen Dateien gültig sind.")